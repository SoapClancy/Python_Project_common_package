from docx import Document
from docx.shared import Cm, Pt
from typing import List
from Writting import new_document
import os
import re
from File_Management.path_and_file_management_Func import try_to_find_file_if_exist_then_delete, \
    list_all_specific_format_files_in_a_path


def put_list_png_file_into_a_docx(png_file_in_a_list: List[str], docx_file_: str):
    document = new_document()
    cols = 2
    rows = int(png_file_in_a_list.__len__() / cols) + 1
    table = document.add_table(rows=rows, cols=cols)
    for i, val in enumerate(png_file_in_a_list):
        row = int(i / 2)
        col = i % 2
        this_cell = table.rows[row].cells[col]
        this_cell.paragraphs[0].add_run().add_picture(val, width=Cm(7.5))
    try_to_find_file_if_exist_then_delete(docx_file_)
    document.save(docx_file_)


def put_all_png_in_a_path_into_a_docx(path_: str, docx_file_: str):
    files = list_all_specific_format_files_in_a_path(path_, 'png')
    put_list_png_file_into_a_docx(files, docx_file_)
