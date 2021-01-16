from typing import List, Union
from pathlib import Path
from Writting import docx_document_template_to_collect_figures
from docx import Document
from docx.shared import Cm, Pt
from File_Management.path_and_file_management_Func import try_to_find_file_if_exist_then_delete, \
    list_all_specific_format_files_in_a_folder_path
from File_Management.load_save_Func import load_pkl_file
from collections import OrderedDict


def put_picture_into_a_docx(picture_in_a_ordered_dict: OrderedDict, docx_file_path: Union[str, Path], cols: int = 2):
    """
    picture_in_a_dict： tree-like DS. The leaf node holds the picture to plot.
    1) Any node except the leaf node is a OrderedDict obj, whose key is a string representing the names of its children,
    and the value is a OrderedDict obj
    2) The leaf node is also a OrderedDict obj, and the key is a string representing the names of its value,
    and the value is a tuple, in the form of "(picture: obj，width: number)"
    """
    assert cols in (1, 2)

    def is_leaf_node(node: OrderedDict) -> bool:
        return True if isinstance(next(iter(node.values())), tuple) else False

    def write_leaf_node(node: OrderedDict):
        assert is_leaf_node(node)
        if cols == 2:
            rows = int(node.__len__()) + 1
        else:  # cols == 1:
            rows = 2 * node.__len__() + 1

        table = document.add_table(rows=rows, cols=cols)
        for i, (key, val) in enumerate(node.items()):
            if cols == 2:
                row = int(i / 2) * 2
                col = i % 2
                this_cell = table.rows[row].cells[col]
                this_cell.paragraphs[0].add_run().add_picture(val[0], width=Cm(7.5))
                table.rows[row + 1].cells[col].paragraphs[0].add_run(key)
            else:
                row = i * 2
                col = 0
                this_cell = table.rows[row].cells[col]
                this_cell.paragraphs[0].add_run().add_picture(val[0], width=Cm(15))
                table.rows[row + 1].cells[col].paragraphs[0].add_run(key)

    def main(node, level=1):
        if not is_leaf_node(node):
            for child_key, child_val in node.items():
                document.add_heading(child_key, level=level)
                main(child_val, level + 1)
        else:
            write_leaf_node(node)

    # Iter, do DFS
    document = docx_document_template_to_collect_figures()
    main(picture_in_a_ordered_dict)

    # Save
    try_to_find_file_if_exist_then_delete(docx_file_path)
    document.save(docx_file_path)


def put_all_png_in_a_path_into_a_docx(path_: Union[str, Path], docx_file_path: Union[str, Path]):
    files = list_all_specific_format_files_in_a_folder_path(path_, 'png', '')
    to_do_ = {f'{i}': (x,) for i, x in enumerate(files)}
    put_picture_into_a_docx(to_do_, docx_file_path)


def put_cached_png_into_a_docx(cached_png: dict,
                               docx_file_path: Union[str, Path],
                               cols: int):
    """
    :param cached_png: 一个dict，key代表标题，value代表图像（BytesIO()和宽度）
    :param docx_file_path
    :param cols
    """
    put_picture_into_a_docx(cached_png, docx_file_path, cols)


def put_cached_png_into_a_docx_new():
    pass


